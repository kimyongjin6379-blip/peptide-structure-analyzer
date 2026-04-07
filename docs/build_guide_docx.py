"""
TOOL_GUIDE.md → Peptide_Structure_Analyzer_Tool_Guide.docx 자동 변환

사용법:
    python docs/build_guide_docx.py

TOOL_GUIDE.md를 수정한 뒤 이 스크립트를 실행하면 docx가 자동 재생성됩니다.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DOCS_DIR = Path(__file__).parent
MD_PATH = DOCS_DIR / 'TOOL_GUIDE.md'
OUT_PATH = DOCS_DIR / 'Peptide_Structure_Analyzer_Tool_Guide.docx'

BRAND_DARK = RGBColor(0x1B, 0x3A, 0x5C)
BRAND_MID = RGBColor(0x4A, 0x6F, 0xA5)
BRAND_GRAY = RGBColor(0x66, 0x66, 0x66)
BRAND_NOTE = RGBColor(0x33, 0x33, 0x99)


# ─── 스타일 초기화 ───────────────────────────────────
def init_styles(doc: Document):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for lv in range(1, 4):
        h = doc.styles[f'Heading {lv}']
        h.font.name = 'Arial'
        h.font.color.rgb = BRAND_DARK

    doc.styles['Heading 1'].font.size = Pt(18)
    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 3'].font.size = Pt(11)


# ─── 표지 ────────────────────────────────────────────
def add_cover(doc: Document, meta: dict):
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get('title', 'Peptide Structure Analyzer'))
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = BRAND_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Tool Guide')
    r.font.size = Pt(20)
    r.font.color.rgb = BRAND_MID

    doc.add_paragraph()

    # meta 줄들 (> 로 시작하는 것들)
    for line in meta.get('subtitle_lines', []):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = BRAND_GRAY

    doc.add_page_break()


# ─── 테이블 생성 ─────────────────────────────────────
def add_md_table(doc: Document, headers: list, rows: list):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = val.strip()
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()


# ─── 코드 블록 ───────────────────────────────────────
def add_code_block(doc: Document, lines: list):
    text = '\n'.join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ─── inline 서식 처리 (bold, code, italic) ───────────
def add_formatted_run(paragraph, text: str, base_size=Pt(10)):
    """마크다운 inline 서식을 docx run으로 변환"""
    # **bold**, `code`, *italic* 패턴 처리
    pattern = r'(\*\*(.+?)\*\*|`(.+?)`|\*(.+?)\*)'
    last = 0
    for m in re.finditer(pattern, text):
        # 이전 일반 텍스트
        if m.start() > last:
            r = paragraph.add_run(text[last:m.start()])
            r.font.size = base_size
        if m.group(2):  # **bold**
            r = paragraph.add_run(m.group(2))
            r.bold = True
            r.font.size = base_size
        elif m.group(3):  # `code`
            r = paragraph.add_run(m.group(3))
            r.font.name = 'Consolas'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x33, 0x00)
        elif m.group(4):  # *italic*
            r = paragraph.add_run(m.group(4))
            r.italic = True
            r.font.size = base_size
        last = m.end()
    # 나머지 텍스트
    if last < len(text):
        r = paragraph.add_run(text[last:])
        r.font.size = base_size


# ─── MD 파싱 & 변환 메인 ─────────────────────────────
def convert_md_to_docx(md_path: Path, out_path: Path):
    md_text = md_path.read_text(encoding='utf-8')
    lines = md_text.split('\n')

    doc = Document()
    init_styles(doc)

    # ── 1단계: 메타 정보 추출 (첫 # 제목 + > 인용문) ──
    meta = {'title': 'Peptide Structure Analyzer', 'subtitle_lines': []}
    content_start = 0

    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith('# ') and not meta.get('_title_found'):
            meta['title'] = stripped[2:].strip()
            meta['_title_found'] = True
            content_start = i + 1
        elif stripped.startswith('> ') and meta.get('_title_found'):
            meta['subtitle_lines'].append(stripped[2:].strip())
            content_start = i + 1
        elif stripped and meta.get('_title_found') and not stripped.startswith('>'):
            break

    add_cover(doc, meta)

    # ── 2단계: 본문 파싱 ──
    i = content_start
    in_code_block = False
    code_lines = []
    in_table = False
    table_headers = []
    table_rows = []
    page_break_pending = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- 코드 블록 ---
        if stripped.startswith('```'):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # 테이블 플러시
                if in_table:
                    add_md_table(doc, table_headers, table_rows)
                    in_table = False
                    table_headers = []
                    table_rows = []
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- 테이블 ---
        if '|' in stripped and stripped.startswith('|') and stripped.endswith('|'):
            cells = [c.strip() for c in stripped.split('|')[1:-1]]

            # 구분선 (|---|---|)
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue

            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                add_md_table(doc, table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []

        # --- 빈 줄 ---
        if not stripped:
            i += 1
            continue

        # --- 수평선 (---) → 페이지 구분 힌트 ---
        if re.match(r'^-{3,}$', stripped):
            # heading 1 앞에만 page break
            page_break_pending = True
            i += 1
            continue

        # --- 헤딩 ---
        heading_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()

            # level 1 (##) 앞에 페이지 나누기
            if level <= 2 and page_break_pending:
                doc.add_page_break()
            page_break_pending = False

            doc.add_heading(text, level=level)
            i += 1
            continue

        page_break_pending = False

        # --- 인용문 (> ) → 정보 박스 ---
        if stripped.startswith('> '):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_run(p, text, base_size=Pt(9))
            for r in p.runs:
                r.font.color.rgb = BRAND_NOTE
                r.italic = True
            i += 1
            continue

        # --- 일반 텍스트 / 리스트 ---
        p = doc.add_paragraph()
        # 리스트 아이템 (- 또는 숫자.)
        list_match = re.match(r'^(\s*)([-•*]|\d+\.)\s+(.+)$', stripped)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            text = list_match.group(3)
            p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.5)
            bullet = list_match.group(2)
            if re.match(r'\d+\.', bullet):
                prefix = f'{bullet} '
            else:
                prefix = '• '
            r = p.add_run(prefix)
            r.font.size = Pt(10)
            add_formatted_run(p, text)
        else:
            add_formatted_run(p, stripped)

        i += 1

    # 마지막 테이블 플러시
    if in_table:
        add_md_table(doc, table_headers, table_rows)

    # ── 저장 ──
    doc.save(str(out_path))
    size_kb = out_path.stat().st_size / 1024
    print(f'Done: {out_path} ({size_kb:.1f} KB)')


if __name__ == '__main__':
    convert_md_to_docx(MD_PATH, OUT_PATH)
